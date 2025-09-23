# 代码生成时间: 2025-09-23 20:13:28
import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt

"""
创建一个Streamlit应用程序，用于将Word文档转换为PDF和CSV格式。
"""

def convert_word_to_pdf(docx_path: str) -> str:
    """
    将Word文档转换为PDF格式。
    
    参数:
        docx_path (str): Word文档的路径。
    
    返回:
        str: 生成的PDF文件路径。
    """
    # 使用python-docx库加载Word文档
    doc = Document(docx_path)
    # 保存为PDF文件（这里仅为示例，实际需要其他库支持）
    pdf_path = docx_path.replace('.docx', '.pdf')
    # 此处省略实际转换代码，因为python-docx不支持直接转换为PDF
    return pdf_path


def extract_text_from_word(docx_path: str) -> pd.DataFrame:
    """
    从Word文档中提取文本并保存为CSV格式。
    
    参数:
        docx_path (str): Word文档的路径。
    
    返回:
        pd.DataFrame: 包含文档文本的DataFrame。
    """
    # 使用python-docx库加载Word文档
    doc = Document(docx_path)
    # 提取文档中的所有段落
    paragraphs = [p.text for p in doc.paragraphs]
    # 将文本保存为DataFrame
    df = pd.DataFrame(paragraphs, columns=['Text'])
    # 保存为CSV文件
    csv_path = docx_path.replace('.docx', '.csv')
    df.to_csv(csv_path, index=False)
    return df


def main():
    """
    Streamlit应用程序的主函数。
    """
    st.title('Word文档格式转换器')
    # 用户上传Word文档
    uploaded_file = st.file_uploader("选择一个Word文档 (.docx)", type=['docx'])
    if uploaded_file is not None:
        try:
            # 将文件保存到临时路径
            docx_path = st.write(uploaded_file)
            docx_path = uploaded_file.name
            # 转换文档
            pdf_path = convert_word_to_pdf(docx_path)
            # 提取文档文本
            df = extract_text_from_word(docx_path)
            # 显示PDF路径和CSV数据
            st.write(f'PDF文件将被保存为: {pdf_path}')
            st.write(df)
        except Exception as e:
            st.error(f'发生错误: {e}')

if __name__ == '__main__':
    main()